import docx
import os
from .utils import URL_PATTERN, clean_urls

def extract_urls_from_docx(filepath):
    urls = []
    try:
        document = docx.Document(filepath)
        # Extract from paragraphs
        for para in document.paragraphs:
            urls.extend(URL_PATTERN.findall(para.text))
        # Extract from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        urls.extend(URL_PATTERN.findall(para.text))
    except Exception as e:
        print(f"--> Error processing DOCX file '{os.path.basename(filepath)}': {e}")
        return []

    return clean_urls(urls)

